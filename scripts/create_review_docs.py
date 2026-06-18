"""
create_review_docs.py — Word-Review-Dokumente für Wissensfreund Mini-Lauf.

Erstellt 6 .docx-Dateien (je Thema), S1+S2+S3 pro Datei.
Liest lektorierte Artikel aus articles/batch_output/lektorat/,
Bilder aus .cache/downloads/{md5(thumb_url)}_300.jpg.
Output: articles/mini_s2_v2/review/
"""

from __future__ import annotations
import hashlib
import io
import json
import os
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent.parent
LEKTORAT_DIR = ROOT / "articles" / "batch_output" / "lektorat"
CACHE_DIR = ROOT / ".cache" / "downloads"
OUT_DIR = ROOT / "articles" / "mini_s2_v2" / "review"

THEMEN = [
    ("Elefant", "elefant"),
    ("Hund", "hund"),
    ("Dinosaurier", "dinosaurier"),
    ("Vulkan", "vulkan"),
    ("Spartacus", "spartacus"),
    ("Zweiter Weltkrieg", "zweiter_weltkrieg"),
]


# ── Farben ────────────────────────────────────────────────────────────────────
CLR_YELLOW   = RGBColor(0xFF, 0xFF, 0x00)
CLR_ORANGE   = RGBColor(0xFF, 0x99, 0x00)
CLR_BLUE     = RGBColor(0x00, 0x70, 0xC0)
CLR_GREEN    = RGBColor(0x00, 0x80, 0x00)
CLR_RED      = RGBColor(0xC0, 0x00, 0x00)
CLR_GRAY     = RGBColor(0x80, 0x80, 0x80)
CLR_LT_GREEN = RGBColor(0xC6, 0xEF, 0xCE)
CLR_LT_YELL  = RGBColor(0xFF, 0xFF, 0xCC)
CLR_LT_ORG   = RGBColor(0xFF, 0xE0, 0xB2)
CLR_LT_BLUE  = RGBColor(0xDD, 0xEE, 0xFF)
CLR_REVIEW   = RGBColor(0xFF, 0xEE, 0xEE)


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────
def img_cache_path(thumb_url: str) -> Path | None:
    if not thumb_url:
        return None
    ck = hashlib.md5(thumb_url.encode()).hexdigest()
    p = CACHE_DIR / f"{ck}_300.jpg"
    return p if p.exists() else None


def set_cell_bg(cell, rgb_hex: str) -> None:
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def set_para_bg(para, rgb_hex: str) -> None:
    """Setzt Hintergrundfarbe eines Absatzes (via shading in pPr)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    pPr.append(shd)


def add_colored_box(doc: Document, text: str, prefix: str,
                    bg_hex: str, title_color: RGBColor) -> None:
    """Fügt eine farbige Box als 1×1-Tabelle ein."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    para = cell.paragraphs[0]
    run = para.add_run(prefix + " ")
    run.bold = True
    run.font.color.rgb = title_color
    run.font.size = Pt(10)
    run2 = para.add_run(text)
    run2.font.size = Pt(10)
    doc.add_paragraph()


def add_stimmt_das_box(doc: Document, text: str, reveal_text: str) -> None:
    """Grüne Box für stimmt_das mit Auflösung."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "E2EFDA")
    para = cell.paragraphs[0]
    run = para.add_run("❓ ")
    run.bold = True
    run.font.color.rgb = CLR_GREEN
    run2 = para.add_run(text)
    run2.font.size = Pt(10)
    if reveal_text:
        p2 = cell.add_paragraph()
        r_label = p2.add_run("Auflösung: ")
        r_label.bold = True
        r_label.font.color.rgb = CLR_GRAY
        r_label.font.size = Pt(9)
        r_rev = p2.add_run(reveal_text)
        r_rev.font.color.rgb = CLR_GRAY
        r_rev.font.size = Pt(9)
    doc.add_paragraph()


def add_warning_banner(doc: Document, review_reason: str) -> None:
    """Roter Banner bei review_flag."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "FFE0E0")
    para = cell.paragraphs[0]
    run = para.add_run(f"⚠ Zur Überprüfung markiert: {review_reason}")
    run.bold = True
    run.font.color.rgb = CLR_RED
    run.font.size = Pt(10)
    doc.add_paragraph()


def add_lektorat_summary(doc: Document, pb: dict) -> None:
    """Lektorat-Zusammenfassung am Artikelende."""
    if not pb:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        set_cell_bg(cell, "F5F5F5")
        para = cell.paragraphs[0]
        para.add_run("Lektorat ausstehend").italic = True
        doc.add_paragraph()
        return

    n_s = pb.get("n_silent", 0)
    n_k = pb.get("n_korrigiert", 0)
    n_p = pb.get("n_pruefen", 0)
    text = pb.get("text", "")

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F0F4FF")

    header_para = cell.paragraphs[0]
    run = header_para.add_run(
        f"LEKTORAT  |  {n_s} SILENT  ·  {n_k} KORRIGIERT  ·  {n_p} PRÜFEN"
    )
    run.bold = True
    run.font.size = Pt(9)

    # Parse pruefbericht.text Zeilen
    current_section = None
    for line in text.split("\n"):
        if line.startswith("## ") or line.startswith("Zusammenfassung:"):
            continue
        if line.startswith("### SILENT"):
            current_section = "silent"
            continue
        if line.startswith("### KORRIGIERT"):
            current_section = "korrigiert"
            p = cell.add_paragraph()
            r = p.add_run("KORRIGIERT")
            r.bold = True
            r.font.color.rgb = CLR_ORANGE
            r.font.size = Pt(8)
            continue
        if line.startswith("### PRÜFEN"):
            current_section = "pruefen"
            p = cell.add_paragraph()
            r = p.add_run("PRÜFEN — Manuelle Kontrolle erforderlich")
            r.bold = True
            r.font.color.rgb = CLR_RED
            r.font.size = Pt(8)
            continue
        if line.startswith("- ") and current_section:
            content = line[2:]
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            if current_section == "silent":
                r = p.add_run(content)
                r.font.color.rgb = CLR_GRAY
                r.font.size = Pt(8)
            elif current_section == "korrigiert":
                # Parse «alt» → «neu» — WP: «...»
                m = re.match(r'«(.+?)» → «(.+?)»(.*)', content)
                if m:
                    r_alt = p.add_run(f"«{m.group(1)}»")
                    r_alt.font.color.rgb = CLR_GRAY
                    r_alt.font.size = Pt(8)
                    r_alt.font.strike = True
                    r_arrow = p.add_run(" → ")
                    r_arrow.font.size = Pt(8)
                    r_neu = p.add_run(f"«{m.group(2)}»")
                    r_neu.font.size = Pt(8)
                    r_neu.bold = True
                    set_para_bg(p, "FFFF99")
                    if m.group(3):
                        r_wp = p.add_run(m.group(3))
                        r_wp.font.size = Pt(7)
                        r_wp.font.color.rgb = CLR_GRAY
                else:
                    r = p.add_run(content)
                    r.font.size = Pt(8)
                    set_para_bg(p, "FFFF99")
            elif current_section == "pruefen":
                r = p.add_run("⚑ " + content)
                r.font.color.rgb = CLR_RED
                r.font.size = Pt(8)

    doc.add_paragraph()


def add_quiz(doc: Document, quiz: dict) -> None:
    """Quiz am Artikelende."""
    if not quiz:
        return
    questions = quiz.get("questions", [])
    if not questions:
        return

    h = doc.add_heading("Quiz", level=2)
    h.style.font.color.rgb = CLR_BLUE

    for i, q in enumerate(questions, 1):
        q_text = q.get("text", q.get("question", ""))
        correct = q.get("correct_key", "")
        options = q.get("options", [])

        q_para = doc.add_paragraph()
        run = q_para.add_run(f"{i}. {q_text}")
        run.bold = True
        run.font.size = Pt(10)

        for opt in options:
            key = opt.get("key", "")
            text = opt.get("text", "")
            is_correct = (key == correct)
            opt_para = doc.add_paragraph()
            opt_para.paragraph_format.left_indent = Inches(0.3)
            if is_correct:
                set_para_bg(opt_para, "C6EFCE")
            run_opt = opt_para.add_run(f"{key})  {text}")
            run_opt.font.size = Pt(9)
            if is_correct:
                run_opt.bold = True
                run_opt.font.color.rgb = CLR_GREEN

    doc.add_paragraph()


def add_image_inline(doc: Document, img_data: dict, width_inches: float,
                     align: str = "center") -> None:
    """Bettet ein Bild ein (oder Platzhalter wenn nicht vorhanden)."""
    thumb_url = img_data.get("thumb_url", "")
    img_path = img_cache_path(thumb_url)
    caption = img_data.get("caption", img_data.get("alt", ""))
    license_str = f"{img_data.get('license', '')} {img_data.get('license_author', '')}".strip()

    para = doc.add_paragraph()
    if align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if img_path:
        run = para.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
    else:
        fname = img_data.get("filename", "Unbekannt")
        run = para.add_run(f"[Bild: {fname[:50]}]")
        run.font.color.rgb = CLR_GRAY
        run.font.size = Pt(8)

    cap_para = doc.add_paragraph()
    if align == "right":
        cap_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = cap_para.add_run(caption)
    run_cap.font.size = Pt(7)
    run_cap.font.color.rgb = CLR_GRAY
    if license_str:
        run_lic = cap_para.add_run(f"  ({license_str})")
        run_lic.font.size = Pt(6)
        run_lic.font.color.rgb = CLR_GRAY


def build_article_section(doc: Document, article: dict, stufe_num: int) -> None:
    """Fügt einen vollständigen Artikel-Abschnitt ins Dokument."""
    meta = article.get("meta", {})
    images = article.get("images", [])
    sections = article.get("sections", [])
    quiz = article.get("quiz", {})
    pb = article.get("pruefbericht", {})

    title = meta.get("title", "?")
    subtitle = meta.get("subtitle", "")
    word_count = meta.get("word_count", 0)
    review_flag = meta.get("review_flag", False)
    review_reason = meta.get("review_reason", "")

    # Abschnitt-Überschrift
    header_para = doc.add_paragraph()
    set_para_bg(header_para, "2F5597")
    r_hdr = header_para.add_run(f"  S{stufe_num}  |  {word_count} Wörter")
    r_hdr.bold = True
    r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_hdr.font.size = Pt(11)

    # review_flag Banner
    if review_flag:
        add_warning_banner(doc, review_reason)

    # Titel + Untertitel
    t_para = doc.add_heading(title, level=1)
    if subtitle:
        sub_para = doc.add_paragraph(subtitle)
        sub_para.runs[0].italic = True
        sub_para.runs[0].font.size = Pt(11)

    # Hero-Bild
    hero_imgs = [img for img in images if img.get("is_hero")]
    if hero_imgs:
        add_image_inline(doc, hero_imgs[0], width_inches=5.8, align="center")

    # Map img_index → Bild für Abschnittsbilder
    img_map: dict[int, dict] = {i: img for i, img in enumerate(images)}
    shown_images: set[int] = {images.index(h) for h in hero_imgs} if hero_imgs else set()

    # Abschnitte
    for sec in sections:
        heading = sec.get("heading", "")
        sentences = sec.get("sentences", [])
        boxes = sec.get("boxes", [])

        if heading:
            doc.add_heading(heading, level=2)

        # Fließtext aufbauen; inline-Bilder beim ersten Satz mit img_index einbetten
        para_text = ""
        section_img_indices: list[int] = []
        for sent in sentences:
            para_text += sent.get("text", "") + " "
            idx = sent.get("img_index", -1)
            if idx >= 0 and idx not in shown_images and idx not in section_img_indices:
                img_d = img_map.get(idx)
                if img_d and img_d.get("ab_stufe", 99) <= stufe_num:
                    section_img_indices.append(idx)

        doc.add_paragraph(para_text.strip())

        # Begleitbilder (max 2 pro Section)
        for idx in section_img_indices[:2]:
            img_d = img_map[idx]
            shown_images.add(idx)
            add_image_inline(doc, img_d, width_inches=2.4, align="right")

        # Boxen
        for box in boxes:
            box_type = box.get("type", "")
            box_text = box.get("text", "")
            if box_text.startswith("BOX:"):
                box_text = box_text[4:].strip()

            if box_type == "wow":
                add_colored_box(doc, box_text, "★", "FFFF99", CLR_ORANGE)
            elif box_type == "fakt":
                add_colored_box(doc, box_text, "ℹ", "DDEEFF", CLR_BLUE)
            elif box_type == "warnung":
                add_colored_box(doc, box_text, "⚠", "FFE0B2", CLR_ORANGE)
            elif box_type == "stimmt_das":
                reveal = box.get("reveal_text", "")
                add_stimmt_das_box(doc, box_text, reveal)
            else:
                add_colored_box(doc, box_text, "▶", "F5F5F5", CLR_GRAY)

    # Quiz
    add_quiz(doc, quiz)

    # Lektorat-Zusammenfassung
    doc.add_heading("Lektorat-Bericht", level=2)
    add_lektorat_summary(doc, pb)


def build_theme_doc(thema_name: str, slug: str,
                    lektorat_dir: Path | None = None) -> Path:
    """Erstellt ein vollständiges Review-Dokument für ein Thema."""
    effective_lektorat_dir = lektorat_dir if lektorat_dir is not None else LEKTORAT_DIR
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    # ── Deckblatt ────────────────────────────────────────────────────────────
    doc.add_heading(f"Wissensfreund — Review: {thema_name}", level=0)
    doc.add_paragraph(f"Datum: {date.today().isoformat()}")
    doc.add_paragraph(f"Thema: {thema_name}  |  Mini-Lauf S1/S2/S3")

    # Artikel laden und Gesamtübersicht
    articles_by_stufe: dict[int, dict] = {}
    total_k = total_p = total_s = 0
    for stufe_num in [1, 2, 3]:
        lekt_path = effective_lektorat_dir / f"lektorat_{slug}_l{stufe_num}.json"
        if not lekt_path.exists():
            continue
        art = json.loads(lekt_path.read_text(encoding="utf-8"))
        articles_by_stufe[stufe_num] = art
        pb = art.get("pruefbericht", {})
        total_s += pb.get("n_silent", 0)
        total_k += pb.get("n_korrigiert", 0)
        total_p += pb.get("n_pruefen", 0)

    n_artikel = len(articles_by_stufe)
    summary_para = doc.add_paragraph()
    summary_para.add_run(
        f"Artikel: {n_artikel}  |  "
        f"SILENT: {total_s}  |  KORRIGIERT: {total_k}  |  PRÜFEN: {total_p}"
    ).bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Farbkodierung:  KORRIGIERT = gelb hinterlegt, Originaltext durchgestrichen  "
        "|  PRÜFEN = roter Hinweis  |  Boxen = farbige Rahmen"
    ).runs[0].font.size = Pt(8)

    # ── Artikel-Abschnitte ────────────────────────────────────────────────────
    for stufe_num in sorted(articles_by_stufe.keys()):
        doc.add_page_break()
        art = articles_by_stufe[stufe_num]
        build_article_section(doc, art, stufe_num)

    # Speichern
    safe_name = thema_name.replace(" ", "")
    out_path = OUT_DIR / f"{safe_name}_Review.docx"
    doc.save(str(out_path))
    return out_path


def main() -> None:
    import argparse
    global OUT_DIR, LEKTORAT_DIR
    sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser()
    parser.add_argument("--out-dir", default=None)
    parser.add_argument("--lektorat-dir", default=None)
    parser.add_argument("--theme-lektorat", action="append", default=[],
                        metavar="SLUG:PATH",
                        help="Per-Thema Lektorat-Dir, z.B. spartacus:/pfad/zu/lektorat")
    args = parser.parse_args()
    if args.out_dir:
        OUT_DIR = Path(args.out_dir).resolve()
    if args.lektorat_dir:
        LEKTORAT_DIR = Path(args.lektorat_dir).resolve()
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Per-Thema Lektorat-Dir-Overrides aufbauen
    theme_lektorat_map: dict[str, Path] = {}
    for entry in args.theme_lektorat:
        slug, _, path_str = entry.partition(":")
        if slug and path_str:
            theme_lektorat_map[slug.strip()] = Path(path_str.strip()).resolve()

    for thema_name, slug in THEMEN:
        print(f"Erstelle {thema_name}_Review.docx ...", end=" ", flush=True)
        try:
            lektorat_dir = theme_lektorat_map.get(slug)
            path = build_theme_doc(thema_name, slug, lektorat_dir=lektorat_dir)
            size_kb = path.stat().st_size // 1024
            print(f"OK ({size_kb} KB) → {path.name}")
        except Exception as e:
            print(f"FEHLER: {e}")
            import traceback
            traceback.print_exc()

    print(f"\nFertig. Alle Dokumente in: {OUT_DIR}")


if __name__ == "__main__":
    main()
