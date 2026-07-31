#!/usr/bin/env python3
"""render_raw_articles_docx.py — liest ROH-Artikel-JSONs (Stage 2, ohne Lektorat)
und rendert sie als lesbares Word-Dokument zum Stil-Check.

Nutzung:
  python -X utf8 scripts/render_raw_articles_docx.py \
      articles/stiltest_20260710/articles/vulkan_l2.json \
      articles/stiltest_20260710/articles/vulkan_l3.json \
      --output "C:/Users/Andreas/Desktop/_review_stiltest_vulkan_20260710.docx"
"""
from __future__ import annotations
import argparse, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

_BOX_LABEL = {
    "wusstest_du": "💡 Wusstest du?",
    "stimmt_das": "🤔 Stimmt das?",
    "schon_gewusst": "💡 Schon gewusst?",
    "merke": "📌 Merke",
}


def _add_article(doc: Document, data: dict, src_name: str):
    meta = data.get("meta", {})
    title = meta.get("title") or src_name
    sub = meta.get("subtitle") or ""
    lvl = meta.get("age_level")
    wc = meta.get("word_count")

    h = doc.add_heading(title, level=1)
    if sub:
        p = doc.add_paragraph()
        r = p.add_run(sub); r.italic = True; r.font.size = Pt(12)
    info = doc.add_paragraph()
    ri = info.add_run(f"[{src_name} · Stufe {lvl} · {wc} Wörter · roh, ohne Lektorat]")
    ri.font.size = Pt(9); ri.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for s in data.get("sections", []):
        doc.add_heading(s.get("heading", ""), level=2)
        text = " ".join((x.get("text") or "").strip() for x in s.get("sentences", []) if x.get("text"))
        if text:
            doc.add_paragraph(text)
        for box in s.get("boxes", []):
            label = _BOX_LABEL.get(box.get("type", ""), box.get("type", "Box"))
            bp = doc.add_paragraph()
            br = bp.add_run(f"{label}: "); br.bold = True
            br.font.color.rgb = RGBColor(0x1F, 0x5C, 0xA8)
            bp.add_run(box.get("text", ""))

    # Quiz kompakt (nur zur Info)
    quiz = data.get("quiz") or []
    if quiz:
        doc.add_heading("Quiz", level=2)
        for i, q in enumerate(quiz, 1):
            if isinstance(q, str):
                doc.add_paragraph(f"{i}. {q}")
                continue
            frage = q.get("frage") or q.get("question") or q.get("text") or ""
            qp = doc.add_paragraph()
            qp.add_run(f"{i}. {frage}").bold = True
            opts = q.get("optionen") or q.get("options") or q.get("antworten") or []
            for o in opts:
                otxt = o if isinstance(o, str) else (o.get("text") or "")
                doc.add_paragraph(otxt, style="List Bullet")

    doc.add_page_break()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("jsons", nargs="+", type=Path)
    ap.add_argument("--output", type=Path, required=True)
    args = ap.parse_args()

    doc = Document()
    intro = doc.add_paragraph()
    ir = intro.add_run("Stil-Check — Roh-Prosa (Stage 2, neuer STIL-&-RHYTHMUS-Block, KEIN Lektorat)")
    ir.bold = True; ir.font.size = Pt(13)
    doc.add_page_break()

    for jf in args.jsons:
        data = json.loads(Path(jf).read_text(encoding="utf-8"))
        _add_article(doc, data, jf.stem)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"OK -> {args.output}")


if __name__ == "__main__":
    main()
