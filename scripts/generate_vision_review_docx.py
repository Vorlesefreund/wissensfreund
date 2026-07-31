#!/usr/bin/env python3
"""generate_vision_review_docx.py — Bild-Kontaktbogen für Vision-Review.

Erzeugt EIN Word-Dokument mit allen Bildern der 3 sensiblen Themen
(Zweiter Weltkrieg + Titanic + Gladiator) samt Gemini-Vision-Urteilen,
damit die Einstufungen (ab_stufe / grenzfall) per Auge geprüft werden können.

Pro Thema eine Überschrift; pro Bild ein Block: 800px-Bild (~10 cm) + kompakte
Metadaten-Tabelle. Grenzfälle werden mit roter "⚠ GRENZFALL"-Zeile + rotem
Zellrahmen hervorgehoben.

Bildquelle: .cache/downloads/{md5(thumb_url)}_800.jpg, Fallback thumb_url-Download.
Output: articles/vision_review_52bilder.docx

Aufruf: python -X utf8 scripts/generate_vision_review_docx.py
"""
from __future__ import annotations

import hashlib
import io
import json
import os
import sys
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT      = Path(__file__).resolve().parent.parent
CACHE_DIR = ROOT / ".cache" / "downloads"
OUT_PATH  = ROOT / "articles" / "vision_review_52bilder.docx"

RUNS = [
    ("Zweiter Weltkrieg", ROOT / "articles" / "vision_ww2_20260625" / "stage1_checkpoint.json"),
    ("Titanic",           ROOT / "articles" / "vision_titanic_20260625" / "stage1_checkpoint.json"),
    ("Gladiator",         ROOT / "articles" / "vision_gladiator_20260625" / "stage1_checkpoint.json"),
]

RED   = RGBColor(0xC0, 0x00, 0x00)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
HEADER_FILL = "E0E0E0"
GF_FILL     = "FDE7E7"   # zartes Rot für Grenzfall-Metazeile


# ── Bild-Cache (wie generate_review_docx.py) ─────────────────────────────────
def get_image_bytes(thumb_url: str, size: str = "800"):
    key = hashlib.md5((thumb_url or "").encode()).hexdigest()
    path = CACHE_DIR / f"{key}_{size}.jpg"
    if path.exists():
        return path.read_bytes(), "cache"
    if thumb_url:
        try:
            r = requests.get(thumb_url, timeout=10,
                             headers={"User-Agent": "Wissensfreund/1.0"})
            if r.status_code == 200:
                return r.content, "download"
        except Exception:
            pass
    return None, "fail"


# ── docx-Low-Level ───────────────────────────────────────────────────────────
def _run(p, text, size=10, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_border(table, color="CCCCCC", sz=8) -> None:
    """Setzt Außen- + Innenrahmen der Tabelle."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def add_meta_row(table, label, value, value_color=None, label_bold=True):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    _run(lc.paragraphs[0], label, size=9, bold=label_bold, color=GRAY)
    _run(vc.paragraphs[0], value or "—", size=9, color=value_color)


# ── Hauptlauf ────────────────────────────────────────────────────────────────
def main() -> None:
    doc = Document()
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width  = Cm(21.0)
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(title, "Vision-Review — Bild-Kontaktbogen (52 Bilder, 3 sensible Themen)",
         size=16, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sub, "Gemini-2.5-Flash-Urteile zur Prüfung · ⚠ GRENZFALL rot markiert",
         size=10, italic=True, color=GRAY)

    total = 0
    total_gf = 0
    missing = 0

    for thema, cp_path in RUNS:
        cp = json.load(open(cp_path, encoding="utf-8"))
        topic_data = cp.get("topics", {})
        # Thema-Überschrift
        doc.add_paragraph()
        h = doc.add_paragraph()
        _run(h, f"━━━  {thema}  ━━━", size=15, bold=True)

        for _t, v in topic_data.items():
            imgs = v.get("images", [])
            n_gf = sum(1 for i in imgs if i.get("grenzfall"))
            cap = doc.add_paragraph()
            _run(cap, f"{len(imgs)} Bilder · {n_gf} Grenzfälle · "
                      f"Companions: {', '.join(v.get('valid_companions', []) or [])}",
                 size=10, italic=True, color=GRAY)

            for im in imgs:
                total += 1
                is_gf = bool(im.get("grenzfall"))
                if is_gf:
                    total_gf += 1
                fname = im.get("filename", "?")

                # Grenzfall-Warnzeile
                if is_gf:
                    w = doc.add_paragraph()
                    _run(w, f"⚠ GRENZFALL — {im.get('grenzfall_grund', '') or 'kein Grund angegeben'}",
                         size=11, bold=True, color=RED)

                # Bild
                data, origin = get_image_bytes(im.get("thumb_url", ""), "800")
                if data:
                    p = doc.add_paragraph()
                    try:
                        p.add_run().add_picture(io.BytesIO(data), width=Cm(10.0))
                    except Exception as e:
                        _run(p, f"[Bild-Einbettung fehlgeschlagen: {fname} — {e}]",
                             size=10, color=RED)
                else:
                    missing += 1
                    p = doc.add_paragraph()
                    _run(p, f"[Bild nicht im Cache: {fname}]", size=10, color=RED)

                # Metadaten-Tabelle
                tbl = doc.add_table(rows=0, cols=2)
                tbl.autofit = False
                tbl.allow_autofit = False
                set_table_border(tbl, color="999999" if is_gf else "CCCCCC")
                add_meta_row(tbl, "Datei", fname)
                add_meta_row(tbl, "Quellartikel", im.get("_source", "?"))
                add_meta_row(tbl, "ab_stufe", f"S{im.get('ab_stufe', '?')}")
                add_meta_row(tbl, "grenzfall", "JA" if is_gf else "nein",
                             value_color=RED if is_gf else GREEN)
                if is_gf:
                    add_meta_row(tbl, "Grund", im.get("grenzfall_grund", "") or "—")
                add_meta_row(tbl, "Relevanz", str(im.get("relevanz", "?")))
                add_meta_row(tbl, "Hero-Kandidat", "JA" if im.get("hero_candidate") else "nein")
                # Spaltenbreiten + Kopf-Schattierung der ersten Zeile
                for row in tbl.rows:
                    row.cells[0].width = Cm(3.5)
                    row.cells[1].width = Cm(13.5)
                if is_gf:
                    for row in tbl.rows:
                        shade_cell(row.cells[0], GF_FILL)

                doc.add_paragraph()  # Abstand zwischen Blöcken

    doc.save(str(OUT_PATH))
    size_kb = OUT_PATH.stat().st_size / 1024
    print(f"✅ Docx erzeugt: {OUT_PATH}")
    print(f"   Bilder: {total} | Grenzfälle: {total_gf} | fehlend(Cache): {missing}")
    print(f"   Größe: {size_kb:.0f} KB")


if __name__ == "__main__":
    main()
