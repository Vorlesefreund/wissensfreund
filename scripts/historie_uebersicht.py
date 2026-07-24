#!/usr/bin/env python3
"""historie_uebersicht.py — alle Fassungen eines Themas als lesbares Docx.

Sammelt jede Artikel-JSON zu den angegebenen Themen aus articles/**, sortiert
chronologisch und schreibt je Thema EIN Docx: vorne eine Uebersichtstabelle
(Datum, Lauf, Stufe, Woerter, Pipeline), danach die Volltexte.

  python -X utf8 scripts/historie_uebersicht.py dinosaurier vulkan
"""
from __future__ import annotations
import argparse, datetime, json, glob, os, re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent

# Alte Stufen-Welt (l1/l2/l3) und neue Inhaltstypen auf eine Spalte bringen.
STUFE_LABEL = {
    "l1": "S1 (alt)", "l2": "S2 (alt)", "l3": "S3 (alt)",
    "hoerspiel": "Hörspiel (4–9)", "erzaehltext": "Erzähltext (10–12)",
}


def stufe_of(basename: str, meta: dict) -> str:
    ct = meta.get("content_type")
    if ct in STUFE_LABEL:
        return STUFE_LABEL[ct]
    m = re.search(r"_(l[123])(?:_|\.)", basename)
    return STUFE_LABEL.get(m.group(1), "?") if m else "?"


def collect(thema: str, highlight: str = "") -> list[dict]:
    out = []
    for f in glob.glob(str(ROOT / "articles" / "**" / "*.json"), recursive=True):
        b = os.path.basename(f)
        if not b.lower().startswith(thema + "_") or b.endswith("_report.json"):
            continue
        try:
            d = json.load(open(f, encoding="utf-8"))
        except Exception:
            continue
        meta = d.get("meta")
        if not isinstance(meta, dict) or not d.get("sections"):
            continue
        lauf = os.path.relpath(os.path.dirname(f), ROOT / "articles").replace("\\", "/")
        out.append({
            "datum": (meta.get("generated_at") or "0000")[:16].replace("T", " "),
            "lauf": lauf,
            "datei": b,
            "pfad": os.path.relpath(f, ROOT).replace("\\", "/"),
            "stufe": stufe_of(b, meta),
            "woerter": meta.get("word_count", 0),
            "pipeline": meta.get("generation_method", "?"),
            "titel": meta.get("subtitle") or meta.get("title", ""),
            "companions": meta.get("grounding_companions") or [],
            # NEU = aus dem aktuellen Lauf (highlight-Substring im Lauf-Pfad)
            "neu": bool(highlight) and highlight in lauf,
            "doc": d,
        })
    out.sort(key=lambda r: (r["datum"], r["stufe"]))
    return out


def add_table(doc: Document, rows: list[dict]) -> None:
    t = doc.add_table(rows=1, cols=7)
    t.style = "Light Grid Accent 1"
    for c, h in zip(t.rows[0].cells,
                    ["", "Datum", "Lauf", "Stufe", "Wörter", "Pipeline", "#"]):
        c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
    for i, r in enumerate(rows, 1):
        cells = t.add_row().cells
        marker = "★ NEU" if r.get("neu") else ""
        for c, v in zip(cells, [marker, r["datum"], r["lauf"], r["stufe"], str(r["woerter"]),
                                r["pipeline"], f"#{i}"]):
            c.text = v
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
                    if marker and v == marker:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)


def add_text(doc: Document, idx: int, r: dict) -> None:
    tag = "★ NEU  " if r.get("neu") else ""
    h = doc.add_heading(f"{tag}#{idx}  {r['datum']}  ·  {r['stufe']}  ·  {r['lauf']}", level=1)
    for run in h.runs:
        run.font.size = Pt(13)
        if tag:
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    p = doc.add_paragraph()
    run = p.add_run(f"{r['titel']}   |   {r['woerter']} Wörter   |   {r['pipeline']}   |   {r['pfad']}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if r["companions"]:
        p = doc.add_paragraph()
        run = p.add_run("Companions: " + ", ".join(str(c) for c in r["companions"]))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for sec in r["doc"].get("sections", []):
        if sec.get("heading"):
            hp = doc.add_heading(sec["heading"], level=2)
            for run in hp.runs:
                run.font.size = Pt(11)
        for s in sec.get("sentences", []):
            para = doc.add_paragraph(s.get("text", ""))
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(10.5)
        for box in sec.get("boxes", []) or []:
            bp = doc.add_paragraph()
            run = bp.add_run(f"[{box.get('type','BOX').upper()}] {box.get('text','')}")
            run.font.size = Pt(9)
            run.italic = True
    doc.add_page_break()


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Alle Fassungen je Thema chronologisch als Docx (alt vs. neu).")
    ap.add_argument("themen", nargs="*", default=["dinosaurier", "vulkan"],
                    help="Themen (Slug, kleingeschrieben)")
    ap.add_argument("--outdir", default="",
                    help="Ziel-Ordner (Default: Desktop/Wissensfreund_Review/<heute>_Historie)")
    ap.add_argument("--stand", default="",
                    help="Datums-Label im Titel/Dateinamen (Default: heute)")
    ap.add_argument("--highlight", default="",
                    help="Lauf-Substring, dessen Fassungen als '★ NEU' markiert werden "
                         "(z. B. der Nachtlauf-Ordner nacht_review2_20260724)")
    ap.add_argument("--fokus", type=int, default=0,
                    help="Auf einen lesbaren Vergleich eindampfen: behalte alle ★NEU-Fassungen "
                         "plus die N jüngsten ALTEN je Thema (0 = alle Fassungen zeigen).")
    args = ap.parse_args()

    themen = [t.lower() for t in args.themen] or ["dinosaurier", "vulkan"]
    stand = args.stand or datetime.date.today().isoformat()
    outdir = (Path(args.outdir) if args.outdir
              else Path.home() / "Desktop" / "Wissensfreund_Review" / f"{stand}_Historie")
    outdir.mkdir(parents=True, exist_ok=True)

    for thema in themen:
        rows = collect(thema, highlight=args.highlight)
        if not rows:
            print(f"{thema}: nichts gefunden")
            continue
        # Fokus: alle ★NEU behalten, dazu nur die N jüngsten ALTEN (chronologisch
        # ist rows aufsteigend sortiert → die letzten N Nicht-NEU sind die jüngsten).
        if args.fokus > 0:
            neu_rows = [r for r in rows if r.get("neu")]
            alt_rows = [r for r in rows if not r.get("neu")][-args.fokus:]
            rows = sorted(neu_rows + alt_rows, key=lambda r: (r["datum"], r["stufe"]))
        n_neu = sum(1 for r in rows if r.get("neu"))
        doc = Document()
        doc.add_heading(f"{thema.capitalize()} — alle Fassungen chronologisch (alt vs. neu)", level=0)
        p = doc.add_paragraph(
            f"{len(rows)} Fassungen"
            + (f" · {n_neu} neu (★) aus dem aktuellen Lauf" if n_neu else "")
            + f" · Stand {stand} · Repo: wissensfreund_repo/articles/")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_table(doc, rows)
        doc.add_page_break()
        for i, r in enumerate(rows, 1):
            add_text(doc, i, r)
        out = outdir / f"{thema.capitalize()}_alt_vs_neu_{stand}.docx"
        doc.save(out)
        print(f"OK: {out}  ({len(rows)} Fassungen, {n_neu} neu, {out.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()
