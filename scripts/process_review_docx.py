#!/usr/bin/env python3
"""process_review_docx.py — Verarbeitet eine vom Menschen kommentierte Review-Docx.

Liest die Kommentare aus der rechten Spalte (3-cm-Kommentarfeld, w:w=1701,
Rahmen 999999) der durchgehend 2-spaltigen Review-Tabelle, ordnet sie den
reviewten Sätzen zu und handelt je nach Ablage-Ordner:

  APPROVED-Ordner  → editorial_approved.json setzen; bei Kommentaren zusätzlich
                     als Änderungswünsche behandeln (wie CHANGES).
  CHANGES-Ordner   → Kommentare = neuer Satztext (es sei denn beginnt mit "!"
                     → Anweisung für Claude Chat, nicht umsetzen); danach neues
                     Docx (vN+1) auf den Desktop legen.

Aufruf:
    python scripts/process_review_docx.py <docx_pfad> [--run-dir <DIR>]

--run-dir: Lauf-Ordner (articles/<run>). Fehlt er, wird er aus dem Docx-Namen
abgeleitet (Konvention: review_<run>_vN.docx).
"""
from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.shared import Cm

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT = Path(__file__).resolve().parent.parent
DESKTOP        = Path(r"C:\Users\Andreas\Desktop")
APPROVED_DIR   = DESKTOP / "wissensfreund_approved"
CHANGES_DIR    = DESKTOP / "wissensfreund_changes"

COMMENT_COL_EMU = int(Cm(3))      # rechte Spalte = 3 cm
COL_TOL_EMU     = int(Cm(0.6))    # Toleranz für Breitenvergleich
MATCH_THRESHOLD = 0.60            # difflib-Mindestähnlichkeit Satz↔Zelle


# ── Hilfen ────────────────────────────────────────────────────────────────────
def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip().lower()


def _atomic_write(path: Path, data: str) -> None:
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(data, encoding="utf-8")
    os.replace(tmp, path)


# ── Kommentar-Extraktion ──────────────────────────────────────────────────────
def extract_comment_rows(docx_path: Path) -> list[dict]:
    """Alle Zeilen der Kommentartabelle: {satz, kommentar}. Nur Zeilen mit Text
    in der rechten 3-cm-Spalte (kommentar nicht leer)."""
    doc = Document(str(docx_path))
    rows: list[dict] = []
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) != 2:
                continue
            rw = cells[1].width
            if rw is None:
                continue
            if abs(int(rw) - COMMENT_COL_EMU) > COL_TOL_EMU:
                continue   # nicht die Kommentarspalte (z. B. Quiz 15,5 cm)
            kommentar = (cells[1].text or "").strip()
            if not kommentar:
                continue
            satz = (cells[0].text or "").strip()
            rows.append({"satz": satz, "kommentar": kommentar})
    return rows


def doc_full_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


# ── Lektorat-Index ────────────────────────────────────────────────────────────
def load_lektorat(run_dir: Path) -> list[dict]:
    """Lädt alle lektorat_*.json eines Laufs. Pro Datei: {path, lk, dirty}."""
    lekt_dir = run_dir / "lektorat"
    out = []
    for fp in sorted(lekt_dir.glob("lektorat_*.json")):
        with open(fp, encoding="utf-8") as fh:
            out.append({"path": fp, "lk": json.load(fh), "dirty": False})
    return out


def build_sentence_index(files: list[dict]) -> list[dict]:
    """Flacher Index aller Sätze: {file, sec_i, sent_i, text, norm, aid}."""
    idx = []
    for entry in files:
        lk = entry["lk"]
        aid = (lk.get("meta", {}) or {}).get("id", entry["path"].stem)
        for si, sec in enumerate(lk.get("sections", []) or []):
            for ji, sent in enumerate(sec.get("sentences", []) or []):
                t = (sent.get("text") or "").strip()
                if t:
                    idx.append({"file": entry, "sec_i": si, "sent_i": ji,
                                "text": t, "norm": _norm(t), "aid": aid})
    return idx


def best_sentence(left_text: str, index: list[dict]):
    """Bestes Satz-Match für den linken Zelltext (exact > substring > fuzzy)."""
    nleft = _norm(left_text)
    if not nleft:
        return None, 0.0
    best, best_r = None, 0.0
    for item in index:
        st = item["norm"]
        if st == nleft:
            return item, 1.0
        if st in nleft or nleft in st:
            r = 0.9
        else:
            r = SequenceMatcher(None, nleft, st).ratio()
        if r > best_r:
            best, best_r = item, r
    return (best, best_r) if best_r >= MATCH_THRESHOLD else (None, best_r)


# ── Marker / Regeneration ─────────────────────────────────────────────────────
def appeared_article_ids(files: list[dict], full_text: str) -> list[str]:
    """Artikel-IDs, deren Header-Signatur '<title>   |   Stufe: S<n>' im Docx steht."""
    out = []
    for entry in files:
        meta = entry["lk"].get("meta", {}) or {}
        aid = meta.get("id", entry["path"].stem)
        title = meta.get("title", "")
        stufe = meta.get("age_level", "")
        sig = f"{title}   |   Stufe: S{stufe}"
        if title and sig in full_text:
            out.append(aid)
    if not out:   # Fallback: Signatur nicht gefunden → alle Lauf-Artikel
        out = [(e["lk"].get("meta", {}) or {}).get("id", e["path"].stem) for e in files]
    return sorted(set(out))


def write_editorial_approved(run_dir: Path, ids: list[str]) -> Path:
    path = run_dir / "editorial_approved.json"
    payload = {"approved": ids,
               "approved_at": datetime.now(timezone.utc).isoformat()}
    _atomic_write(path, json.dumps(payload, ensure_ascii=False, indent=2))
    return path


def next_version(docx_path: Path) -> int:
    m = re.search(r"_v(\d+)", docx_path.stem)
    return (int(m.group(1)) if m else 1) + 1


def regenerate_docx(run_dir: Path, docx_path: Path) -> tuple[Path | None, str]:
    out = DESKTOP / f"review_{run_dir.name}_v{next_version(docx_path)}.docx"
    cmd = [sys.executable, str(ROOT / "scripts" / "generate_review_docx.py"),
           str(run_dir), "--output", str(out)]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")
    except Exception as e:
        return None, f"Aufruf fehlgeschlagen: {e}"
    if r.returncode != 0:
        return None, (r.stderr or r.stdout or "unbekannter Fehler").strip()[-400:]
    return out, "ok"


# ── Hauptlogik ────────────────────────────────────────────────────────────────
def resolve_run_dir(docx_path: Path, explicit: str | None) -> Path:
    if explicit:
        p = Path(explicit)
        return (p if p.is_absolute() else ROOT / p).resolve()
    # Konvention: review_<run>_vN.docx
    stem = docx_path.stem
    name = re.sub(r"^review_", "", stem)
    name = re.sub(r"_v\d+$", "", name)
    cand = (ROOT / "articles" / name).resolve()
    if not (cand / "lektorat").is_dir():
        print(f"FEHLER: run_dir nicht ableitbar aus '{docx_path.name}' "
              f"(probiert: {cand}). Bitte --run-dir angeben.", file=sys.stderr)
        sys.exit(2)
    return cand


def detect_mode(docx_path: Path) -> str:
    parent = docx_path.resolve().parent
    s = str(parent).lower()
    if APPROVED_DIR.name.lower() in s:
        return "approved"
    if CHANGES_DIR.name.lower() in s:
        return "changes"
    return "auto"


def main() -> int:
    ap = argparse.ArgumentParser(description="Verarbeitet kommentierte Review-Docx.")
    ap.add_argument("docx", type=Path)
    ap.add_argument("--run-dir", default=None)
    args = ap.parse_args()

    docx_path = args.docx.resolve()
    if not docx_path.is_file():
        print(f"FEHLER: Docx nicht gefunden: {docx_path}", file=sys.stderr)
        return 2

    run_dir = resolve_run_dir(docx_path, args.run_dir)
    mode = detect_mode(docx_path)

    rows = extract_comment_rows(docx_path)
    has_comments = len(rows) > 0
    if mode == "auto":   # außerhalb der bekannten Ordner: aus Kommentaren ableiten
        mode = "changes" if has_comments else "approved"

    do_approve = (mode == "approved")
    do_apply   = (mode == "changes") or (mode == "approved" and has_comments)

    print(f"Docx:    {docx_path.name}")
    print(f"Lauf:    {run_dir.name}   Modus: {mode}")
    print(f"Zeilen mit Kommentar: {len(rows)}\n")

    files = load_lektorat(run_dir)
    if not files:
        print(f"FEHLER: keine lektorat_*.json in {run_dir/'lektorat'}", file=sys.stderr)
        return 2
    index = build_sentence_index(files)

    claude_items: list[str] = []
    n_ersetzt = n_claude = n_unmatched = 0

    for r in rows:
        satz, kommentar = r["satz"], r["kommentar"]
        kurz = (satz[:60] + "…") if len(satz) > 60 else satz

        if kommentar.startswith("!"):
            anweisung = kommentar[1:].strip()
            claude_items.append(f"SATZ: {satz}\nANWEISUNG: {anweisung}\n")
            n_claude += 1
            print(f"  [claude_chat] {kurz!r}  ⇒  {anweisung!r}")
            continue

        item, ratio = best_sentence(satz, index)
        if item is None:
            claude_items.append(f"UNMATCHED-SATZ: {satz}\nKOMMENTAR: {kommentar}\n")
            n_unmatched += 1
            print(f"  [unmatched r={ratio:.2f}] {kurz!r}  ⇒  {kommentar!r}")
            continue

        if do_apply:
            sec = item["file"]["lk"]["sections"][item["sec_i"]]
            sec["sentences"][item["sent_i"]]["text"] = kommentar
            item["file"]["dirty"] = True
            n_ersetzt += 1
            print(f"  [ersetzt r={ratio:.2f}] ({item['aid']}) {kurz!r}  ⇒  {kommentar!r}")
        else:   # reiner Approved-Fall mit (theoretisch) Kommentar, kein Apply
            print(f"  [approved-notiz r={ratio:.2f}] {kurz!r}  ⇒  {kommentar!r}")

    # Geänderte Lektorate speichern
    saved = 0
    if do_apply:
        for entry in files:
            if entry["dirty"]:
                _atomic_write(entry["path"],
                              json.dumps(entry["lk"], ensure_ascii=False, indent=2))
                saved += 1

    full_text = doc_full_text(docx_path)

    # editorial_approved.json
    approved_path = None
    if do_approve:
        ids = appeared_article_ids(files, full_text)
        approved_path = write_editorial_approved(run_dir, ids)

    # changes_for_claude.txt
    claude_path = None
    if claude_items:
        claude_path = run_dir / "changes_for_claude.txt"
        header = (f"# Änderungswünsche für Claude Chat — {docx_path.name}\n"
                  f"# erzeugt {datetime.now(timezone.utc).isoformat()}\n\n")
        _atomic_write(claude_path, header + "\n".join(claude_items))

    # Docx-Regeneration (nur wenn Änderungen angewandt wurden)
    new_docx = regen_msg = None
    if do_apply and saved:
        new_docx, regen_msg = regenerate_docx(run_dir, docx_path)

    # ── Bericht ───────────────────────────────────────────────────────────────
    print()
    print(f"ersetzt: {n_ersetzt} | claude_chat: {n_claude} | unmatched: {n_unmatched} "
          f"| Lektorate gespeichert: {saved}")
    if approved_path:
        print(f"editorial_approved.json gesetzt → {approved_path} "
              f"({len(json.loads(approved_path.read_text(encoding='utf-8'))['approved'])} IDs)")
    if claude_path:
        print(f"changes_for_claude.txt erzeugt → {claude_path} ({n_claude + n_unmatched} Einträge)")
    if new_docx:
        print(f"neues Review-Docx → {new_docx}")
    elif regen_msg and regen_msg != "ok":
        print(f"WARNUNG Docx-Regeneration: {regen_msg}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
