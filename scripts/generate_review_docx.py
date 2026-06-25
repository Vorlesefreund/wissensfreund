#!/usr/bin/env python3
"""generate_review_docx.py — Word-Review-Dokument für einen Pipeline-Lauf.

Durchgehend 2-spaltig: jedes Element (Überschrift, Satz, Bild, Box,
Lektorat-Marking) ist eine Zeile einer 2-Spalten-Tabelle —
links Inhalt (14 cm), rechts leeres Kommentarfeld (3 cm, sichtbarer Rahmen).
Header + Hero + Quiz bleiben volle Breite.

Aufruf:
    python scripts/generate_review_docx.py articles/test_erde_regenwald_wal \
        --output articles/test_erde_regenwald_wal/review_erw_20260624_v3.docx
    [--themen erde]

Bildquelle: .cache/downloads/{md5(thumb_url)}_{800}.jpg, Fallback thumb_url-Download.
"""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

CACHE_DIR = ".cache/downloads"
EXCLUDE_BOX_CONTENT = ["Osedax"]

COL_HEADER  = "E0E0E0"
COL_WOW     = "E8F4FD"
COL_WARNUNG = "FFF9C4"
COL_STIMMT  = "E8F5E9"
COL_FAKT    = "F5F5F5"
COL_CORRECT = "C8E6C9"
COL_SECHEAD = "EFEFEF"

GRAY  = RGBColor(0x77, 0x77, 0x77)
RED   = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x1B, 0x5E, 0x20)

_stats: list[dict] = []
_warnings: list[str] = []


# ── Filter ──────────────────────────────────────────────────────────────────--
def _excluded(text: str) -> bool:
    t = text or ""
    return any(s in t for s in EXCLUDE_BOX_CONTENT)


def _strip_box_prefix(s: str) -> str:
    return re.sub(r"^BOX\[[^\]]*\]:\s*", "", s or "")


# ── Bild-Cache ──────────────────────────────────────────────────────────────--
def get_image_bytes(thumb_url: str, size: str = "800"):
    key = hashlib.md5(thumb_url.encode()).hexdigest()
    path = os.path.join(CACHE_DIR, f"{key}_{size}.jpg")
    if os.path.exists(path):
        with open(path, "rb") as fh:
            return fh.read(), "cache"
    try:
        r = requests.get(thumb_url, timeout=10, headers={"User-Agent": "Wissensfreund/1.0"})
        if r.status_code == 200:
            return r.content, "download"
    except Exception:
        pass
    return None, "fail"


# ── docx-Low-Level ───────────────────────────────────────────────────────────-
def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC", sz=8) -> None:
    """Setzt alle vier Zellrahmen (single). sz in Achtel-Punkt (8 ≈ 1pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_col_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _run(p, text, size=11, bold=False, italic=False, color=None, strike=False):
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.strike = strike
    if color is not None:
        r.font.color.rgb = color
    return r


def cell_text(cell, text, size=10, bold=False, italic=False, color=None):
    _run(cell.paragraphs[0], text, size=size, bold=bold, italic=italic, color=color)


# ── 2-Spalten-Tabelle ────────────────────────────────────────────────────────-
def add_two_col_table(doc):
    tbl = doc.add_table(rows=0, cols=2)   # kein Style → randlos; Rahmen pro Zelle
    tbl.autofit = False
    tbl.allow_autofit = False
    return tbl


def _two_col_row(tbl):
    """Neue Zeile; links Inhalt (14cm), rechts leeres Kommentarfeld (3cm, Rahmen)."""
    row = tbl.add_row()
    left, right = row.cells[0], row.cells[1]
    left.width = Cm(14.0)
    right.width = Cm(3.0)
    _set_cell_border(left, color="CCCCCC")
    _set_cell_border(right, color="999999")   # Kommentarfeld deutlicher sichtbar
    return left, right


# ── Bild in Zelle ─────────────────────────────────────────────────────────────
def _img_caption_text(im) -> str:
    lic = im.get("license", "")
    aut = im.get("license_author", "")
    txt = im.get("caption") or ""
    if lic:
        txt += f"  ({lic}{' — ' + aut if aut else ''})"
    return txt


def render_image_cell(cell, im, width_cm=6.0) -> str:
    data, origin = get_image_bytes(im["thumb_url"], "800")
    p = cell.paragraphs[0]
    if data is None:
        _run(p, f"[Bild nicht verfügbar: {im.get('filename','?')}]", 9, italic=True, color=GRAY)
        _warnings.append(f"Bild fehlt: {im.get('filename')}")
        return "fail"
    try:
        p.add_run().add_picture(io.BytesIO(data), width=Cm(width_cm))
    except Exception as e:
        _run(p, f"[Bild-Fehler: {im.get('filename','?')} — {e}]", 9, italic=True, color=GRAY)
        _warnings.append(f"Bild-Insert-Fehler {im.get('filename')}: {e}")
        return "fail"
    _run(cell.add_paragraph(), _img_caption_text(im), 9, italic=True, color=GRAY)
    return origin


# ── Bausteine (volle Breite) ──────────────────────────────────────────────────
def render_image_fullwidth(doc, im, width_cm=14.0) -> str:
    data, origin = get_image_bytes(im["thumb_url"], "800")
    if data is None:
        _run(doc.add_paragraph(), f"[Bild nicht verfügbar: {im.get('filename','?')}]", 9, italic=True, color=GRAY)
        _warnings.append(f"Hero fehlt: {im.get('filename')}")
        return "fail"
    try:
        doc.add_picture(io.BytesIO(data), width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        _run(doc.add_paragraph(), f"[Bild-Fehler: {im.get('filename','?')} — {e}]", 9, italic=True, color=GRAY)
        return "fail"
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cap, _img_caption_text(im), 9, italic=True, color=GRAY)
    return origin


def add_header_block(doc, meta: dict) -> None:
    stufe = meta.get("age_level", "?")
    wc    = meta.get("word_count", "?")
    flag  = "✓" if meta.get("review_flag") else "—"
    line = (f"{meta.get('emoji','')}  {meta.get('title', meta.get('id',''))}   |   "
            f"Stufe: S{stufe}   |   Wörter: {wc}   |   review_flag: {flag}")
    tbl = doc.add_table(rows=1, cols=1)
    set_col_widths(tbl, [17.0])
    cell = tbl.cell(0, 0)
    shade_cell(cell, COL_HEADER)
    _run(cell.paragraphs[0], line, size=14, bold=True)


def add_hero_image(doc, images, stufe_num):
    hero_pos = None
    for pos, im in enumerate(images):
        if im.get("is_hero") and im.get("ab_stufe", 1) <= stufe_num:
            hero_pos = pos
            break
    if hero_pos is None:
        for pos, im in enumerate(images):
            if im.get("is_hero"):
                hero_pos = pos
                break
    if hero_pos is None:
        return "none", None
    return render_image_fullwidth(doc, images[hero_pos], 14.0), hero_pos


def _finding_lookups(findings):
    by_corrected, by_original = {}, {}
    for f in findings:
        claim = _strip_box_prefix(f.get("claim_original") or "").strip()
        korr  = _strip_box_prefix(f.get("korrektur_neu") or "").strip()
        if korr:
            by_corrected[korr] = f
        if claim:
            by_original.setdefault(claim, f)
    return by_corrected, by_original


def _render_tracked_change(left_cell, f: dict) -> None:
    p = left_cell.paragraphs[0]
    _run(p, f"[{f.get('verdikt','')}] ", 11, italic=True, color=GRAY)
    orig = _strip_box_prefix(f.get("claim_original") or "")
    new  = _strip_box_prefix(f.get("korrektur_neu") or f.get("korrektur_vorschlag") or "")
    _run(p, orig, 11, color=RED, strike=True)
    if new:
        _run(p, "  ", 11)
        _run(p, new, 11, bold=True, color=GREEN)
    verdikt = f.get("verdikt", "")
    if verdikt == "PRÜFEN":
        problem = f.get("problem") or ""
        beleg   = f.get("beleg") or ""
        if problem:
            pp = left_cell.add_paragraph()
            _run(pp, "Quelle: " + problem[:300] + ("…" if len(problem) > 300 else ""),
                 size=9, italic=True, color=GRAY)
        if beleg:
            pb = left_cell.add_paragraph()
            _run(pb, "Beleg: " + beleg[:300] + ("…" if len(beleg) > 300 else ""),
                 size=9, italic=True, color=GRAY)


def add_box_cell(left, box) -> bool:
    if _excluded(box.get("text", "")) or _excluded(box.get("reveal_text", "")):
        return False
    btype = (box.get("type") or "").lower()
    color = {"wow": COL_WOW, "warnung": COL_WARNUNG,
             "stimmt_das": COL_STIMMT, "fakt": COL_FAKT}.get(btype, COL_FAKT)
    label = {"wow": "WOW", "warnung": "WARNUNG",
             "stimmt_das": "STIMMT DAS?", "fakt": "FAKT"}.get(btype, btype.upper())
    shade_cell(left, color)
    p = left.paragraphs[0]
    _run(p, f"[{label}] ", 10, bold=True, italic=True)
    _run(p, box.get("text", ""), 10, italic=True)
    if btype == "stimmt_das" and box.get("reveal_text"):
        _run(left.add_paragraph(), "→ " + box["reveal_text"], 10, italic=True, color=GRAY)
    return True


def add_body_2col(doc, tbl, sections, images, findings, hero_pos):
    img_map = {pos: im for pos, im in enumerate(images)}
    shown = set()
    if hero_pos is not None:
        shown.add(hero_pos)
    by_corrected, by_original = _finding_lookups(findings)
    handled = set()
    n_cache = n_dl = 0

    for sec in sections:
        left, _ = _two_col_row(tbl)
        shade_cell(left, COL_SECHEAD)
        _run(left.paragraphs[0], sec.get("heading", ""), 12, bold=True)

        for sent in sec.get("sentences", []):
            t = (sent.get("text") or "").strip()
            f = by_corrected.get(t) or by_original.get(t)
            left, _ = _two_col_row(tbl)
            if f is not None and id(f) not in handled:
                _render_tracked_change(left, f)
                handled.add(id(f))
            else:
                _run(left.paragraphs[0], t, 11)
            idx = sent.get("img_index")
            if idx is not None and idx in img_map and idx not in shown:
                licell, _ = _two_col_row(tbl)
                origin = render_image_cell(licell, img_map[idx], 6.0)
                shown.add(idx)
                n_cache += origin == "cache"
                n_dl += origin == "download"

        for box in sec.get("boxes", []) or []:
            left, _ = _two_col_row(tbl)
            if not add_box_cell(left, box):
                # ausgeschlossene Box → Zeile wieder entfernen
                tbl._tbl.remove(left._tc.getparent())

    leftover = [(pos, im) for pos, im in img_map.items() if pos not in shown]
    if leftover:
        left, _ = _two_col_row(tbl)
        shade_cell(left, COL_SECHEAD)
        _run(left.paragraphs[0], "Weitere Bilder", 12, bold=True)
        for pos, im in leftover:
            licell, _ = _two_col_row(tbl)
            origin = render_image_cell(licell, im, 6.0)
            shown.add(pos)
            n_cache += origin == "cache"
            n_dl += origin == "download"
    return n_cache, n_dl, len(shown)


def add_quiz(doc, quiz) -> int:
    questions = (quiz or {}).get("questions", []) or []
    if not questions:
        return 0
    h = doc.add_heading(level=2)
    _run(h, "Quiz", size=12, bold=True)
    for q in questions:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        set_col_widths(tbl, [1.5, 15.5])
        qcell = tbl.cell(0, 0).merge(tbl.cell(0, 1))
        cell_text(qcell, q.get("text", ""), size=11, bold=True)
        correct = q.get("correct_key")
        for opt in q.get("options", []):
            row = tbl.add_row().cells
            key = opt.get("key", "")
            is_corr = (key == correct)
            cell_text(row[0], ("✓ " if is_corr else "") + key, size=10, bold=is_corr)
            cell_text(row[1], opt.get("text", ""), size=10)
            if is_corr:
                shade_cell(row[0], COL_CORRECT)
                shade_cell(row[1], COL_CORRECT)
        doc.add_paragraph()
    return len(questions)


# ── Hauptlauf ─────────────────────────────────────────────────────────────────
def _sort_key(path: Path):
    stem = path.name[len("lektorat_"):-len(".json")]
    if "_l" in stem:
        theme, _, st = stem.rpartition("_l")
        try:
            return (theme, int(st))
        except ValueError:
            return (stem, 0)
    return (stem, 0)


def build(run_dir: Path, output: Path, themen) -> None:
    lekt_dir = run_dir / "lektorat"
    files = sorted(lekt_dir.glob("lektorat_*.json"), key=_sort_key)
    if themen:
        tset = {t.lower() for t in themen}
        files = [f for f in files if _sort_key(f)[0].lower() in tset]
    if not files:
        print(f"FEHLER: keine Lektorat-JSONs in {lekt_dir} (Filter: {themen})", file=sys.stderr)
        sys.exit(2)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t, "Wissensfreund — Review", size=24, bold=True)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sub, f"Lauf: {run_dir.name}", size=14)
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(d, f"Erstellt: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", 11, color=GRAY)
    nn = doc.add_paragraph(); nn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(nn, f"{len(files)} Artikel · rechte Spalte = Kommentarfeld", 11, color=GRAY)

    for fp in files:
        with open(fp, encoding="utf-8") as fh:
            lk = json.load(fh)
        doc.add_page_break()
        meta = lk.get("meta", {})
        images = lk.get("images", []) or []
        stufe_num = meta.get("age_level", 1)
        findings = [f for f in (lk.get("pruefbericht", {}) or {}).get("findings", []) or []
                    if not _excluded(f.get("claim_original", ""))]

        add_header_block(doc, meta)
        hero_origin, hero_pos = add_hero_image(doc, images, stufe_num)
        tbl = add_two_col_table(doc)
        nc, nd, nshown = add_body_2col(doc, tbl, lk.get("sections", []) or [],
                                       images, findings, hero_pos)
        nq = add_quiz(doc, lk.get("quiz", {}) or {})

        hc = nc + (1 if hero_origin == "cache" else 0)
        hd = nd + (1 if hero_origin == "download" else 0)
        _stats.append({"id": meta.get("id", fp.stem), "img_cache": hc, "img_dl": hd,
                       "img_shown": nshown, "findings": len(findings), "quiz": nq,
                       "images": len(images)})

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> int:
    ap = argparse.ArgumentParser(description="Word-Review-Dokument (durchgehend 2-spaltig).")
    ap.add_argument("run_dir", type=Path)
    ap.add_argument("--output", type=Path, default=None)
    ap.add_argument("--themen", nargs="+", default=None)
    args = ap.parse_args()

    run_dir = args.run_dir.resolve()
    output = Path(args.output or (run_dir / "review.docx")).resolve()

    t0 = time.time()
    build(run_dir, output, args.themen)
    dt = time.time() - t0

    size = os.path.getsize(output)
    print(f"\nOK: {output}  ({size/1024:.1f} KB, {dt:.1f}s)")
    print(f"{'Artikel':12} {'img(cache/dl)':>14} {'gezeigt':>8} {'findings':>9} {'quiz':>5} {'images[]':>9}")
    for s in _stats:
        img = f"{s['img_cache']}/{s['img_dl']}"
        print(f"{s['id']:12} {img:>14} {s['img_shown']:>8} "
              f"{s['findings']:>9} {s['quiz']:>5} {s['images']:>9}")
    if _warnings:
        print(f"\n{len(_warnings)} Warnung(en):")
        for w in _warnings[:20]:
            print(f"  - {w}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
